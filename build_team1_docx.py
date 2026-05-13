"""Build the Team 1 testing deliverable as a Word .docx file."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def set_table_rtl(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    bidi = OxmlElement("w:bidiVisual")
    tblPr.append(bidi)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    set_rtl(h)
    for run in h.runs:
        run.font.name = "Arial"
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:cs"), "Arial")
    return h


def add_para(doc, text, bold=False, rtl=True, mono=False):
    p = doc.add_paragraph()
    if rtl:
        set_rtl(p)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Consolas" if mono else "Arial"
    if mono:
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), "Consolas")
        rFonts.set(qn("w:hAnsi"), "Consolas")
    run.bold = bold
    return p


def add_code_block(doc, code):
    """Add a code block. Each line in its own paragraph with monospace font."""
    for line in code.splitlines():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), "Consolas")
        rFonts.set(qn("w:hAnsi"), "Consolas")
        # light gray background via shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2F2F2")
        pPr.append(shd)


def add_table(doc, headers, rows, rtl=True):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    if rtl:
        set_table_rtl(table)
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        if rtl:
            set_rtl(p)
        run = p.add_run(h)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10)
    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            if rtl:
                set_rtl(p)
            run = p.add_run(val)
            run.font.name = "Arial"
            run.font.size = Pt(10)
    return table


def add_bullet(doc, text, rtl=True):
    p = doc.add_paragraph(style="List Bullet")
    if rtl:
        set_rtl(p)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)


# ─────────────────────────────────────────────────────────────────────────
# Build the document
# ─────────────────────────────────────────────────────────────────────────

doc = Document()

# Set base font and RTL doc-level direction
style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)

# Title
title = doc.add_heading("דרישות בדיקות ואיכות — צוות 1", level=0)
set_rtl(title)

add_para(doc, "חברי הצוות: נדב אילוז, נדב מגן", bold=True)
add_para(doc, "חלק בפיתוח: מסך התחברות, רישום עובדים, ניהול לקוחות", bold=True)
add_para(doc, "שפת פיתוח: C# (Windows Forms) | שמירת נתונים: Excel דרך ClosedXML", bold=True)

# ─── 1. CFG ─────────────────────────────────────────────────────────────
add_heading(doc, "1. שני CFG (Control Flow Graphs)", level=1)

# 1.1
add_heading(doc, "1.1 פונקציה ראשונה: IsValidUsername", level=2)

code1 = '''public static bool IsValidUsername(string username, out string error)
{
    error = "";                                              // N1
    if (string.IsNullOrWhiteSpace(username))                 // N2
    {
        error = "Username is required.";
        return false;                                        // N3 -> EXIT
    }
    if (username.Length < 6 || username.Length > 8)          // N4
    {
        error = "Username must be 6-8 characters.";
        return false;                                        // N5 -> EXIT
    }
    int digitCount = 0;                                      // N6
    foreach (char c in username)                             // N7 (loop)
    {
        if (char.IsDigit(c))                                 // N8
            digitCount++;                                    // N9
        else if (!IsEnglishLetter(c))                        // N10
        {
            error = "Username may contain only English letters and digits.";
            return false;                                    // N11 -> EXIT
        }
    }
    if (digitCount > 2)                                      // N12
    {
        error = "Username may contain at most 2 digits.";
        return false;                                        // N13 -> EXIT
    }
    return true;                                             // N14 -> EXIT
}'''
add_code_block(doc, code1)

add_para(doc, "צמתים (Nodes):", bold=True)
add_table(doc, ["צומת", "תיאור"], [
    ["N1", 'אתחול error = ""'],
    ["N2", "תנאי: string.IsNullOrWhiteSpace(username)"],
    ["N3", "החזרת false (שם משתמש ריק)"],
    ["N4", "תנאי: אורך שם המשתמש מחוץ לטווח 6–8"],
    ["N5", "החזרת false (אורך לא תקין)"],
    ["N6", "אתחול digitCount = 0"],
    ["N7", "ראש לולאה: foreach c in username"],
    ["N8", "תנאי: char.IsDigit(c)"],
    ["N9", "digitCount++"],
    ["N10", "תנאי: !IsEnglishLetter(c)"],
    ["N11", "החזרת false (תו לא חוקי)"],
    ["N12", "תנאי: digitCount > 2"],
    ["N13", "החזרת false (יותר מ-2 ספרות)"],
    ["N14", "החזרת true (תקין)"],
    ["EXIT", "סיום הפונקציה"],
])

add_para(doc, "קשתות (Edges):", bold=True)
edges1 = '''N1 → N2
N2 → N3   (true)
N2 → N4   (false)
N3 → EXIT
N4 → N5   (true)
N4 → N6   (false)
N5 → EXIT
N6 → N7
N7 → N8   (יש עוד תו)
N7 → N12  (סיום הלולאה)
N8 → N9   (true)
N8 → N10  (false)
N9 → N7
N10 → N11 (true)
N10 → N7  (false)
N11 → EXIT
N12 → N13 (true)
N12 → N14 (false)
N13 → EXIT
N14 → EXIT'''
add_code_block(doc, edges1)

add_para(doc, "Cyclomatic Complexity = E − N + 2 = 19 − 14 + 2 = 7")
add_para(doc, "הערה: ציירו את הדיאגרמה ב-draw.io / PowerPoint לפי טבלת הצמתים והקשתות, וצרפו כתמונה.")

# 1.2
add_heading(doc, "1.2 פונקציה שנייה: IsValidPassword", level=2)

code2 = '''public static bool IsValidPassword(string password, out string error)
{
    error = "";                                              // M1
    if (string.IsNullOrEmpty(password))                      // M2
    {
        error = "Password is required.";
        return false;                                        // M3 -> EXIT
    }
    if (password.Length < 8 || password.Length > 10)         // M4
    {
        error = "Password must be 8-10 characters.";
        return false;                                        // M5 -> EXIT
    }
    bool hasLetter = false;                                  // M6
    bool hasDigit = false;
    bool hasSpecial = false;
    foreach (char c in password)                             // M7 (loop)
    {
        if (char.IsLetter(c)) hasLetter = true;              // M8 -> M9
        else if (char.IsDigit(c)) hasDigit = true;           // M10 -> M11
        else if (c == '!' || c == '#' || c == '$')           // M12
            hasSpecial = true;                               // M13
    }
    if (!hasLetter || !hasDigit || !hasSpecial)              // M14
    {
        error = "Password must contain ...";
        return false;                                        // M15 -> EXIT
    }
    return true;                                             // M16 -> EXIT
}'''
add_code_block(doc, code2)

add_para(doc, "צמתים:", bold=True)
add_table(doc, ["צומת", "תיאור"], [
    ["M1", "אתחול error"],
    ["M2", "תנאי: סיסמה ריקה"],
    ["M3", "החזרת false"],
    ["M4", "תנאי: אורך סיסמה מחוץ ל-8–10"],
    ["M5", "החזרת false"],
    ["M6", "אתחול 3 דגלים: hasLetter / hasDigit / hasSpecial"],
    ["M7", "ראש לולאה: foreach c"],
    ["M8", "תנאי: char.IsLetter(c)"],
    ["M9", "hasLetter = true"],
    ["M10", "תנאי: char.IsDigit(c)"],
    ["M11", "hasDigit = true"],
    ["M12", "תנאי: c ∈ {!, #, $}"],
    ["M13", "hasSpecial = true"],
    ["M14", "תנאי: לא קיים אחד מ-3 הדגלים"],
    ["M15", "החזרת false"],
    ["M16", "החזרת true"],
])

add_para(doc, "קשתות:", bold=True)
edges2 = '''M1→M2 ; M2→M3(t) ; M2→M4(f) ; M3→EXIT
M4→M5(t) ; M4→M6(f) ; M5→EXIT
M6→M7
M7→M8 (יש תו) ; M7→M14 (סיום)
M8→M9(t) ; M8→M10(f) ; M9→M7
M10→M11(t) ; M10→M12(f) ; M11→M7
M12→M13(t) ; M12→M7(f) ; M13→M7
M14→M15(t) ; M14→M16(f)
M15→EXIT ; M16→EXIT'''
add_code_block(doc, edges2)

add_para(doc, "Cyclomatic Complexity = E − N + 2 = 22 − 16 + 2 = 8")

# ─── 2. User Stories ────────────────────────────────────────────────────
add_heading(doc, "2. שלושה סיפורי משתמש (User Stories)", level=1)

add_heading(doc, "US-1 — רישום עובד חדש", level=2)
add_para(doc, "בתור מזכירה חדשה במרפאה, אני רוצה להירשם למערכת בעזרת שם משתמש, סיסמה, מספר עובד, מייל, ת\"ז ותפקיד, כדי שאוכל להתחבר בהמשך ולנהל לקוחות.")
add_para(doc, "קריטריונים לקבלה:", bold=True)
add_bullet(doc, "שם משתמש 6–8 תווים, עד 2 ספרות, השאר אותיות באנגלית.")
add_bullet(doc, "סיסמה 8–10 תווים, עם לפחות אות, ספרה ותו מיוחד (!, #, $).")
add_bullet(doc, "מספר עובד = 4 ספרות בדיוק.")
add_bullet(doc, "ת\"ז 9 ספרות עם ספרת ביקורת תקינה (אלגוריתם ת\"ז ישראלית).")
add_bullet(doc, "בלחיצה על Register העובד נשמר בקובץ ה-Excel.")

add_heading(doc, "US-2 — התחברות", level=2)
add_para(doc, "בתור וטרינר רשום, אני רוצה להתחבר עם שם משתמש וסיסמה, כדי לקבל גישה לתפריט הראשי בהתאם לתפקידי.")
add_para(doc, "קריטריונים:", bold=True)
add_bullet(doc, "צירוף תקין → פתיחת MainForm עם שם המשתמש בכותרת.")
add_bullet(doc, "צירוף לא תקין → הודעת שגיאה והשארה במסך ההתחברות.")

add_heading(doc, "US-3 — חיפוש לקוח לפי ת\"ז או טלפון", level=2)
add_para(doc, "בתור מזכירה, אני רוצה להזין ת\"ז או טלפון ולבחור את אופן החיפוש, כדי למצוא במהירות לקוח קיים ולראות את החיות שבבעלותו.")
add_para(doc, "קריטריונים:", bold=True)
add_bullet(doc, "בחירה National ID → סינון לפי הכלת המחרוזת בשדה הת\"ז.")
add_bullet(doc, "בחירה Phone → סינון לפי הכלת המחרוזת בשדה הטלפון.")
add_bullet(doc, "לחיצה על שורה מציגה ברשימה הימנית את החיות שבבעלות הלקוח.")

# ─── 3. Test Cases ──────────────────────────────────────────────────────
add_heading(doc, "3. ארבעה מקרי בדיקה (משני סיפורי משתמש שונים)", level=1)

add_heading(doc, "מקרי בדיקה הנגזרים מ-US-1 (רישום עובד)", level=2)
add_table(doc, ["TC ID", "תיאור", "קלט", "תוצאה צפויה", "תוצאה בפועל"], [
    ["TC-1.1", "רישום תקין של מזכירה",
     "user=reception, pass=Clinic$99, emp=1001, name=Dana Levi, mail=d@x.co.il, id=123456782, role=Secretary",
     "חלון Registered מופיע, העובדת מופיעה ב-Employees ב-Excel", "כצפוי"],
    ["TC-1.2", "סיסמה ללא תו מיוחד",
     "user=vetuser, pass=Pass1234 (אורך 8, חסר תו מיוחד)",
     "הודעת שגיאה: Password must contain at least one letter, one digit and one special character", "כצפוי"],
])

add_heading(doc, "מקרי בדיקה הנגזרים מ-US-3 (חיפוש לקוח)", level=2)
add_table(doc, ["TC ID", "תיאור", "קלט", "תוצאה צפויה", "תוצאה בפועל"], [
    ["TC-3.1", "חיפוש לפי ת\"ז קיימת",
     'קיים לקוח עם ת"ז 123456782. מקלידים 123456782 ובוחרים National ID',
     "הטבלה מציגה שורה אחת – הלקוח", "כצפוי"],
    ["TC-3.2", "חיפוש לפי טלפון שאינו קיים",
     "קיים לקוח עם טלפון 0501234567. מקלידים 0509999999, Phone",
     "הטבלה ריקה, אין שגיאה", "כצפוי"],
])

# ─── 4. Functional + GUI ────────────────────────────────────────────────
add_heading(doc, "4. שתי בדיקות פונקציונאליות + שתי בדיקות GUI", level=1)

add_heading(doc, "4.1 בדיקות פונקציונאליות (Functional)", level=2)

add_para(doc, "FT-1: ולידציית סיסמה", bold=True)
add_para(doc, "מה נבדק: לוגיקת IsValidPassword (כללי 8–10 תווים, אות + ספרה + תו מיוחד).")
add_para(doc, "איך: הזנו 5 סיסמאות שונות בטופס Register (קצרה, ארוכה מדי, ללא ספרה, ללא תו מיוחד, תקינה).")
add_para(doc, "תוצאה: 4 הראשונות נדחו עם הודעת שגיאה ספציפית, החמישית עברה בהצלחה. ✓")

add_para(doc, "FT-2: שיוך חיות ללקוח", bold=True)
add_para(doc, "מה נבדק: CustomersForm מציגה ברשימה הימנית את כל החיות עם OwnerNationalId השווה ל-NationalId של הלקוח שנבחר.")
add_para(doc, "איך: יצרנו לקוח, הוספנו 2 חיות שלו ועוד 1 של לקוח אחר, ובחרנו את הלקוח בטבלה.")
add_para(doc, "תוצאה: הופיעו רק 2 החיות שלו. ✓")

add_heading(doc, "4.2 בדיקות GUI", level=2)

add_para(doc, "GT-1: השדה Password מוסתר", bold=True)
add_para(doc, "מה נבדק: ב-LoginForm וב-RegisterForm תווי הסיסמה מוצגים כנקודות (UseSystemPasswordChar = true).")
add_para(doc, "איך: פתחנו כל אחד מהמסכים, הקלדנו abc123!.")
add_para(doc, "תוצאה: הוצגו •••••••. ✓")

add_para(doc, "GT-2: כפתור ← Main Menu סוגר את CustomersForm ומחזיר ל-MainForm", bold=True)
add_para(doc, "מה נבדק: הפונקציונליות של ניווט חזרה.")
add_para(doc, "איך: התחברנו, פתחנו את מסך הלקוחות, לחצנו על הכפתור.")
add_para(doc, "תוצאה: החלון נסגר, MainForm מקבל focus. ✓")

# ─── 5. Test Scenario ───────────────────────────────────────────────────
add_heading(doc, "5. תרחיש בדיקה + 2 מקרי בדיקה + תסריטי בדיקה", level=1)

add_heading(doc, "תרחיש בדיקה (Test Scenario)", level=2)
add_para(doc, "TS-LOGIN: אימות זהות בכניסה למערכת.", bold=True)
add_para(doc, "המשתמש מזין שם משתמש וסיסמה ב-LoginForm. המערכת בודקת מול רשימת Employees שב-Excel: אם יש התאמה — נפתח MainForm; אחרת מוצגת הודעת שגיאה.")

add_heading(doc, "מקרי בדיקה הנגזרים מהתרחיש", level=2)
add_table(doc, ["TC ID", "תיאור", "תוצאה צפויה"], [
    ["TC-LOGIN-1", "התחברות עם פרטים נכונים של עובד שנרשם", "פתיחת MainForm עם שם המשתמש בכותרת"],
    ["TC-LOGIN-2", "התחברות עם סיסמה שגויה", "הודעה Invalid username or password, שדה הסיסמה מתאפס ומקבל focus"],
])

add_heading(doc, "תסריט TC-LOGIN-1", level=2)
add_table(doc, ["שלב", "פעולה", "קלט", "תוצאה צפויה"], [
    ["1", "הפעלת ClinicVets.exe", "—", "מסך התחברות נפתח"],
    ["2", "הקלדה בשדה Username", "vetadmin", "האותיות מופיעות בשדה"],
    ["3", "הקלדה בשדה Password", "Clinic$99", "מוצגות 9 נקודות"],
    ["4", "לחיצה על Login", "—", 'המסך נסגר, MainForm נפתח עם הכותרת "Welcome, Dana Levi"'],
])

add_heading(doc, "תסריט TC-LOGIN-2", level=2)
add_table(doc, ["שלב", "פעולה", "קלט", "תוצאה צפויה"], [
    ["1", "הפעלת ClinicVets.exe", "—", "מסך התחברות נפתח"],
    ["2", "הקלדה בשדה Username", "vetadmin", "האותיות מופיעות"],
    ["3", "הקלדה בשדה Password", "WrongPass1!", "מוצגות 11 נקודות"],
    ["4", "לחיצה על Login", "—", "חלון שגיאה Invalid username or password"],
    ["5", "לחיצה על OK בחלון השגיאה", "—", "חלון השגיאה נסגר, שדה Password ריק וב-focus"],
])

# ─── Submission notes ──────────────────────────────────────────────────
add_heading(doc, "איך להגיש למורה", level=1)
add_bullet(doc, "ציירו את שני ה-CFG ב-draw.io / PowerPoint על פי רשימות הצמתים והקשתות, ושמרו כתמונה בתוך המסמך.")
add_bullet(doc, "צרפו צילומי מסך של ההרצות (Register תקין, שגיאת סיסמה, התחברות, חיפוש לקוח, הצגת חיות הלקוח, ניווט בכפתור Main Menu).")
add_bullet(doc, "שמרו כ-PDF והגישו ביחד עם קובץ ה-.exe וקישור לתיקייה המשותפת.")

# Save
out = "/Users/magen/Desktop/projectC/Testing-Team1.docx"
doc.save(out)
print(f"Wrote: {out}")
